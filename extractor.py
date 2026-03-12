"""
extractor.py
Extraction de texte depuis PDF, DOCX, TXT
avec fallback OCR pour les documents scannés.
"""

import os
import re
import logging
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# EXTRACTION PDF
# ─────────────────────────────────────────────

def extract_pdf(file_path: str) -> str:
    """Extrait le texte d'un PDF. Fallback OCR si le texte est trop court."""
    text = ""

    # Tentative avec pdfplumber (meilleur pour texte natif)
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            text = "\n\n".join(pages)
        logger.info(f"pdfplumber: {len(text)} caractères extraits")
    except ImportError:
        logger.warning("pdfplumber non disponible, tentative PyMuPDF")
    except Exception as e:
        logger.warning(f"pdfplumber échoué: {e}")

    # Fallback PyMuPDF si pdfplumber ne donne rien
    if len(text.strip()) < 100:
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            pages = [page.get_text() for page in doc]
            text = "\n\n".join(pages)
            doc.close()
            logger.info(f"PyMuPDF: {len(text)} caractères extraits")
        except ImportError:
            logger.warning("PyMuPDF non disponible")
        except Exception as e:
            logger.warning(f"PyMuPDF échoué: {e}")

    # Fallback OCR si le document semble être scanné
    if len(text.strip()) < 100:
        logger.info("Document potentiellement scanné, tentative OCR...")
        text = extract_pdf_ocr(file_path)

    return text


def extract_pdf_ocr(file_path: str) -> str:
    """OCR via Tesseract pour les PDFs scannés."""
    try:
        import pytesseract
        from pdf2image import convert_from_path

        images = convert_from_path(file_path, dpi=300)
        texts = []
        for img in images:
            # Langues FR + EN
            page_text = pytesseract.image_to_string(img, lang="fra+eng")
            texts.append(page_text)
        return "\n\n".join(texts)
    except ImportError:
        logger.warning("pytesseract ou pdf2image non disponible. OCR désactivé.")
        return ""
    except Exception as e:
        logger.error(f"OCR échoué: {e}")
        return ""


# ─────────────────────────────────────────────
# EXTRACTION DOCX
# ─────────────────────────────────────────────

def extract_docx(file_path: str) -> str:
    """Extrait le texte d'un fichier DOCX."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extraction des tableaux
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n[TABLEAUX]\n" + "\n".join(table_texts)

        return full_text
    except ImportError:
        logger.warning("python-docx non disponible")
        return ""
    except Exception as e:
        logger.error(f"Erreur DOCX: {e}")
        return ""


# ─────────────────────────────────────────────
# EXTRACTION TXT
# ─────────────────────────────────────────────

def extract_txt(file_path: str) -> str:
    """Lit un fichier texte brut."""
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    return ""


# ─────────────────────────────────────────────
# DISPATCHER PRINCIPAL
# ─────────────────────────────────────────────

def extract_text(file_path: str) -> Tuple[str, str]:
    """
    Extrait le texte selon le type de fichier.
    Retourne (texte_brut, type_fichier).
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_path), "pdf"
    elif ext in [".docx", ".doc"]:
        return extract_docx(file_path), "docx"
    elif ext == ".txt":
        return extract_txt(file_path), "txt"
    else:
        raise ValueError(f"Format non supporté: {ext}")


# ─────────────────────────────────────────────
# NETTOYAGE DU TEXTE
# ─────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Nettoie le texte extrait :
    - Supprime les headers/footers répétitifs
    - Normalise les espaces et sauts de ligne
    - Supprime les caractères parasites
    """
    if not text:
        return ""

    # Supprimer les numéros de page isolés
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)

    # Normaliser les sauts de ligne multiples
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Normaliser les espaces multiples
    text = re.sub(r"[ \t]{2,}", " ", text)

    # Supprimer les lignes ne contenant que des tirets ou points
    text = re.sub(r"^\s*[-–—=_\.]{3,}\s*$", "", text, flags=re.MULTILINE)

    return text.strip()
